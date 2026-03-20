"""RAG (Retrieval Augmented Generation) module for rule-based review."""

from __future__ import annotations

import os
from pathlib import Path
from typing import Optional

from src.core.models import RagConfig


class RAGManager:
    """Manager for RAG-based rule retrieval."""
    
    def __init__(self, config: RagConfig, api_key: Optional[str] = None):
        self.config = config
        self.api_key = api_key
        self._vector_store = None
        self._embeddings = None
    
    def initialize(self) -> None:
        """Initialize RAG components."""
        if not self.config.enabled:
            return
        
        self._init_embeddings()
        self._load_or_build_index()
    
    def _init_embeddings(self) -> None:
        """Initialize embedding model based on configuration.
        
        Supports:
        - OpenAI models (text-embedding-ada-002, etc.)
        - Local HuggingFace models (bge-small-zh-v1.5, all-MiniLM-L6-v2, etc.)
        """
        if self.config.local_model_path:
            try:
                from langchain_community.embeddings import HuggingFaceEmbeddings
                
                model_path = self._find_model_path(self.config.local_model_path)
                
                self._embeddings = HuggingFaceEmbeddings(
                    model_name=model_path,
                )
                return
            except ImportError:
                raise ImportError(
                    f"Failed to load local model from {self.config.local_model_path}"
                )
        
        embedding_model = self.config.embedding_model
        
        if embedding_model.startswith("text-embedding-"):
            try:
                from langchain_openai import OpenAIEmbeddings
                
                kwargs = {"model": embedding_model}
                if self.api_key:
                    kwargs["api_key"] = self.api_key
                
                self._embeddings = OpenAIEmbeddings(**kwargs)
                return
            except ImportError:
                pass
        
        use_local = os.environ.get("USE_LOCAL_EMBEDDING", "true").lower() == "true"
        
        if use_local:
            try:
                from langchain_community.embeddings import HuggingFaceEmbeddings
                
                model_name = embedding_model if embedding_model else "sentence-transformers/all-MiniLM-L6-v2"
                self._embeddings = HuggingFaceEmbeddings(
                    model_name=model_name,
                )
                return
            except ImportError:
                pass
        
        try:
            from langchain_openai import OpenAIEmbeddings
            
            kwargs = {
                "model": self.config.embedding_model,
            }
            if self.api_key:
                kwargs["api_key"] = self.api_key
            
            self._embeddings = OpenAIEmbeddings(**kwargs)
        except ImportError:
            try:
                from langchain_community.embeddings import HuggingFaceEmbeddings
                
                model_name = "sentence-transformers/all-MiniLM-L6-v2"
                self._embeddings = HuggingFaceEmbeddings(
                    model_name=model_name,
                )
            except ImportError:
                raise ImportError(
                    "Please install langchain-community for embeddings"
                )
    
    def _find_model_path(self, base_path: str) -> str:
        """Find the actual model folder from cache directory.
        
        HuggingFace cache structure:
        models--xxx--yyy/
        └── snapshots/
            └── <hash>/
                ├── config.json
                └── model files
        """
        import json
        
        base_path = os.path.expanduser(base_path)
        
        config_file = os.path.join(base_path, "config.json")
        if os.path.exists(config_file):
            return base_path
        
        snapshots_path = os.path.join(base_path, "snapshots")
        if os.path.exists(snapshots_path):
            for snapshot in os.listdir(snapshots_path):
                snapshot_path = os.path.join(snapshots_path, snapshot)
                if os.path.isdir(snapshot_path):
                    config_file = os.path.join(snapshot_path, "config.json")
                    if os.path.exists(config_file):
                        try:
                            with open(config_file, 'r') as f:
                                json.load(f)
                            return snapshot_path
                        except:
                            pass
        
        raise ValueError(
            f"Could not find model in {base_path}. "
            "Please ensure the folder contains config.json and model files."
        )
    
    def _load_or_build_index(self) -> None:
        """Load existing index or build new one."""
        if os.path.exists(self.config.vector_store_dir):
            self._load_index()
        else:
            self._vector_store = None
    
    def _load_index(self) -> None:
        """Load existing vector store."""
        try:
            from langchain_community.vectorstores import Chroma
            
            self._vector_store = Chroma(
                persist_directory=self.config.vector_store_dir,
                embedding_function=self._embeddings,
            )
        except Exception:
            self._vector_store = None
    
    def build_index(self, rule_documents: list[str]) -> None:
        """Build vector index from rule documents."""
        if not self.config.enabled or not rule_documents:
            return
        
        if self._embeddings is None:
            self._init_embeddings()
        
        try:
            try:
                from langchain.text_splitter import MarkdownTextSplitter
            except ImportError:
                try:
                    from langchain_community.text_splitters import MarkdownTextSplitter
                except ImportError:
                    from langchain_community.text_splitter import MarkdownTextSplitter
            
            from langchain_community.vectorstores import Chroma
            
            all_chunks = []
            for doc in rule_documents:
                splitter = MarkdownTextSplitter(
                    chunk_size=self.config.chunk_size,
                    chunk_overlap=self.config.chunk_overlap,
                )
                chunks = splitter.split_text(doc)
                all_chunks.extend(chunks)
            
            if all_chunks:
                self._vector_store = Chroma.from_texts(
                    texts=all_chunks,
                    embedding=self._embeddings,
                    persist_directory=self.config.vector_store_dir,
                )
                self._vector_store.persist()
        except ImportError as e:
            raise ImportError(
                f"Missing required package: {e}. "
                "Please install: pip install langchain langchain-community"
            )
        except Exception as e:
            raise RuntimeError(f"Failed to build RAG index: {e}")
    
    def _simple_chunk_text(self, text: str) -> list[str]:
        """Simple text chunking without langchain dependency."""
        import re
        
        chunks = []
        lines = text.split('\n')
        current_chunk = []
        current_size = 0
        
        for line in lines:
            line_size = len(line)
            
            if current_size + line_size > self.config.chunk_size and current_chunk:
                chunks.append('\n'.join(current_chunk))
                current_chunk = []
                current_size = 0
            
            current_chunk.append(line)
            current_size += line_size + 1
        
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
        
        return chunks
    
    def retrieve_relevant_rules(
        self,
        code_context: str,
        top_k: Optional[int] = None,
    ) -> list[str]:
        """Retrieve relevant rules for given code context."""
        if not self.config.enabled or self._vector_store is None:
            return []
        
        k = top_k or self.config.top_k
        
        try:
            docs = self._vector_store.similarity_search(
                code_context,
                k=k,
            )
            return [doc.page_content for doc in docs]
        except Exception:
            return []
    
    def retrieve_rules_with_intent(
        self,
        file_path: str,
        source_code: str,
        top_k: Optional[int] = None,
    ) -> list[str]:
        """Retrieve relevant rules using AST-based code intent extraction.
        
        This method extracts the natural language intent from code using AST,
        then uses that intent for vector search instead of raw code.
        """
        if not self.config.enabled or self._vector_store is None:
            return []
        
        try:
            from src.rag.intent import extract_code_intent
            
            intent = extract_code_intent(file_path, source_code)
            
            k = top_k or self.config.top_k
            
            docs = self._vector_store.similarity_search(
                intent,
                k=k,
            )
            
            return [doc.page_content for doc in docs]
        except Exception:
            return []
    
    def is_ready(self) -> bool:
        """Check if RAG is ready for use."""
        return self.config.enabled and self._vector_store is not None
    
    def cleanup(self) -> None:
        """Cleanup resources."""
        self._vector_store = None
        self._embeddings = None


def load_rule_documents(paths: list[str]) -> list[str]:
    """Load rule documents from paths.
    
    Supported formats:
    - Markdown (.md)
    - Word (.docx)
    - PDF (.pdf)
    - Plain text (.txt)
    """
    documents = []
    
    for path in paths:
        path_obj = Path(path)
        
        if not path_obj.exists():
            continue
        
        if path_obj.is_file():
            content = _load_single_file(path_obj)
            if content:
                documents.append(content)
        
        elif path_obj.is_dir():
            for ext in ['*.md', '*.docx', '*.pdf', '*.txt']:
                for file_path in path_obj.rglob(ext):
                    content = _load_single_file(file_path)
                    if content:
                        documents.append(content)
    
    return documents


def _load_single_file(file_path: Path) -> Optional[str]:
    """Load content from a single file based on its extension."""
    ext = file_path.suffix.lower()
    
    try:
        if ext == '.md':
            return _load_text_file(file_path)
        elif ext == '.txt':
            return _load_text_file(file_path)
        elif ext == '.docx':
            return _load_docx_file(file_path)
        elif ext == '.pdf':
            return _load_pdf_file(file_path)
    except Exception:
        pass
    
    return None


def _load_text_file(file_path: Path) -> Optional[str]:
    """Load plain text or markdown file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            if content.strip():
                return content
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='gbk') as f:
                content = f.read()
                if content.strip():
                    return content
        except Exception:
            pass
    except Exception:
        pass
    
    return None


def _load_docx_file(file_path: Path) -> Optional[str]:
    """Load Word document (.docx)."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        content = '\n'.join(paragraphs)
        if content.strip():
            return content
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx files. "
            "Install with: pip install python-docx"
        )
    except Exception:
        pass
    
    return None


def _load_pdf_file(file_path: Path) -> Optional[str]:
    """Load PDF document."""
    try:
        import pypdf
        
        reader = pypdf.PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)
        
        content = '\n'.join(text_parts)
        if content.strip():
            return content
    except ImportError:
        raise ImportError(
            "pypdf is required for .pdf files. "
            "Install with: pip install pypdf"
        )
    except Exception:
        pass
    
    return None
